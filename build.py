#!/usr/bin/env python3
"""
build.py — Job Description Library builder
==========================================

Reads every .docx in ./jds, extracts the standardized fields from the
Ant International JD template (v2.x) and writes ./data.js, which the page
(index.html) loads to render the searchable library.

This runs automatically via GitHub Actions whenever a file in /jds changes,
so nobody edits the page by hand — the manager just drops files in /jds.

Run locally:  python build.py
Requires:     pip install python-docx
"""

import glob
import json
import os
import re
import sys

try:
    from docx import Document
except ImportError:
    sys.exit("Missing dependency. Run:  pip install python-docx")

JDS_DIR = "jds"
OUTPUT = "data.js"

# Section headings used by the template, normalized to lowercase.
KNOWN_HEADINGS = {
    "about ant international": "about_ant",
    "our mission": "mission",
    "our vision": "vision",
    "position overview": "position_overview",
    "about the role": "about_role",
    "key responsibilities": "responsibilities",
    "qualifications & experience": "qualifications",
    "qualifications and experience": "qualifications",
    "data privacy notice": "privacy",
    "required": "required",
    "preferred": "preferred",
}


def classify(text):
    """Return a section key for a heading paragraph, else None (= body text)."""
    t = text.lower().strip()
    if t in KNOWN_HEADINGS:
        return KNOWN_HEADINGS[t]
    if t.startswith("about "):          # "About WorldFirst", "About Antom", ...
        return "about_bu"
    if t.startswith("template version"):
        return "footer"
    return None


def slugify(text):
    s = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return s or "jd"


def parse_overview_table(doc):
    """The 'Position Overview' table maps Field -> Details (Job Title, etc.)."""
    fields = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0] and cells[0].lower() != "field":
                fields.setdefault(cells[0], cells[1])
    return fields


def short_location(full):
    """'Sydney, New South Wales, Australia' -> 'Sydney'; 'Hong Kong (Hybrid)' -> 'Hong Kong'."""
    return full.split(",")[0].split("(")[0].strip()


def parse_jd(path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    sections, bu_name, footer = {}, None, ""
    current = "_title"
    for text in paras:
        key = classify(text)
        if key == "about_bu":
            bu_name = text[len("About "):].strip()
            current = "about_bu"
            sections.setdefault(current, [])
            continue
        if key == "footer":
            footer = text
            current = "_footer"
            continue
        if key is not None:
            current = key
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(text)

    ov = parse_overview_table(doc)
    position = ov.get("Job Title") or (sections.get("_title", [""])[0].split(",")[0].strip())
    location_full = ov.get("Location", "")
    fname = os.path.basename(path)

    # Business Unit: prefer the "About <BU>" heading; fall back to the title line.
    bu = bu_name
    if not bu:
        title_parts = [p.strip() for p in sections.get("_title", [""])[0].split(",")]
        bu = title_parts[1] if len(title_parts) >= 2 else ""

    tmpl = re.search(r"Template Version:\s*([\d.]+)", footer)

    return {
        "id": slugify(os.path.splitext(fname)[0]),
        "position": position,
        "bu": bu,
        "local": short_location(location_full),
        "localFull": location_full or short_location(location_full),
        "department": ov.get("Department", ""),
        "reportsTo": ov.get("Reports To", ""),
        "employmentType": ov.get("Employment Type", ""),
        "template": ("v" + tmpl.group(1)) if tmpl else "",
        "file": "jds/" + fname,
        "aboutBu": " ".join(sections.get("about_bu", [])),
        "aboutRole": " ".join(sections.get("about_role", [])),
        "responsibilities": sections.get("responsibilities", []),
        "required": sections.get("required", []),
        "preferred": sections.get("preferred", []),
    }


def main():
    files = sorted(glob.glob(os.path.join(JDS_DIR, "*.docx")))
    files = [f for f in files if not os.path.basename(f).startswith("~$")]  # skip Word temp files
    if not files:
        print(f"No .docx files found in ./{JDS_DIR}")

    records, errors = [], []
    for path in files:
        try:
            rec = parse_jd(path)
            if not rec["position"]:
                raise ValueError("could not detect Job Title")
            records.append(rec)
            print(f"  ✓ {os.path.basename(path)}  ->  {rec['position']} · {rec['bu']} · {rec['local']}")
        except Exception as e:  # noqa: BLE001 — keep building even if one file is malformed
            errors.append((path, str(e)))
            print(f"  ✗ {os.path.basename(path)}  ({e})")

    records.sort(key=lambda r: (r["position"], r["bu"]))

    banner = ("// AUTO-GENERATED by build.py — do not edit by hand.\n"
              "// Source of truth = the .docx files in /jds.\n")
    with open(OUTPUT, "w", encoding="utf-8") as fh:
        fh.write(banner)
        fh.write("window.JDS = ")
        json.dump(records, fh, ensure_ascii=False, indent=2)
        fh.write(";\n")

    print(f"\nWrote {OUTPUT} with {len(records)} job description(s).")
    if errors:
        print(f"{len(errors)} file(s) skipped due to errors.")


if __name__ == "__main__":
    main()
